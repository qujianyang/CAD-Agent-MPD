from io import BytesIO

from docx import Document
from docx.oxml.ns import qn

from catalog import AUTO_SELECT_CATALOGS, select_and_analyze
from physics_engine import ShockEnv
from shock_analysis_context import build_selection_snapshot
from supplier_enquiry_report import (
    CONTENT_WIDTH_DXA,
    generate_supplier_enquiry_pack,
    supplier_enquiry_pack_bytes,
)


def _selection_snapshot():
    report, candidates = select_and_analyze(
        850.0,
        6,
        4,
        shock_env=ShockEnv(Ao_G=20.0, to_s=0.011, GT_limit_G=10.0),
        catalog=AUTO_SELECT_CATALOGS,
    )
    return build_selection_snapshot(
        report,
        candidates,
        analysis_key=("supplier-pack", 850.0, 6, 4),
    )


def _document_text(doc):
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    table_cells = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def test_supplier_pack_contains_authoritative_result_and_review_boundary():
    snapshot = _selection_snapshot()

    doc = generate_supplier_enquiry_pack(
        snapshot,
        project_reference="Project Atlas",
        equipment_reference="Server Rack SR-01",
        prepared_by="Test Engineer",
        supplier="Helical",
        layout_box_mm=(600.0, 800.0, 1200.0),
        clearances_mm=(25.0, 25.0, 25.0),
        wall_face="back",
        report_date="2026-07-26",
    )
    text = _document_text(doc)

    assert "Preliminary Supplier Enquiry Pack" in text
    assert snapshot.analysis_id in text
    assert snapshot.selected_part in text
    assert "850 kg" in text
    assert "20 G" in text
    assert "11 ms" in text
    assert "Four-case calculation results" in text
    assert "Requirement completeness" in text
    assert "Evidence status and traceability" in text
    assert "Supplier nonlinear simulation" in text
    assert "Physical laboratory test" in text
    assert "Road-trial and physical-evidence record" in text
    assert "Supplier confirmation" in text
    assert "Pending" in text
    assert "does not replace supplier confirmation" in text
    for load_case in snapshot.load_cases:
        assert load_case.name in text


def test_supplier_pack_records_extended_requirements_and_evidence_status():
    doc = generate_supplier_enquiry_pack(
        _selection_snapshot(),
        layout_box_mm=(2240.0, 800.0, 1685.0),
        clearances_mm=(30.0, 35.0, 40.0),
        wall_face="back",
        cg_mm=(1120.0, 400.0, 950.0),
        wall_stabilizer_height_mm=1450.0,
        vibration_profile="MIL-STD-810H 514.8C-VII Category 4",
        vibration_duration_min=40.0,
        operating_state="Powered and operating during transport",
        interface_requirements="M8 inserts; supplier to confirm torque",
        environment_requirements="Marine corrosion environment",
        road_trial_status="Required, not started",
    )
    text = _document_text(doc)

    assert "1120 / 400 / 950 mm" in text
    assert "1.079 m/s" in text
    assert "MIL-STD-810H 514.8C-VII Category 4" in text
    assert "40 min" in text
    assert "Powered and operating during transport" in text
    assert "M8 inserts; supplier to confirm torque" in text
    assert "Marine corrosion environment" in text
    assert "Required, not started" in text
    assert "TO BE CONFIRMED" not in "\n".join(
        row.cells[1].text
        for table in doc.tables
        if table.rows[0].cells[0].text == "Requirement"
        for row in table.rows[1:]
    )


def test_supplier_pack_exposes_missing_inputs_and_model_boundaries():
    doc = generate_supplier_enquiry_pack(_selection_snapshot())
    text = _document_text(doc)

    assert "TO BE CONFIRMED" in text
    assert "ASSUMED" in text
    assert "effective configuration value is not a physical mount count" in text
    assert "Shock acceptance and random-vibration duration compliance" in text
    assert "not transferable approval" in text
    assert "Qualification only when supported by an applicable test report" in text


def test_supplier_pack_includes_requested_mount_counts_and_coordinates():
    snapshot = _selection_snapshot()
    doc = generate_supplier_enquiry_pack(
        snapshot,
        layout_box_mm=(600.0, 800.0, 1200.0),
        clearances_mm=(25.0, 25.0, 25.0),
        wall_face="back",
    )
    text = _document_text(doc)

    assert "6 bottom + 4 wall isolators" in text
    assert "Conceptual bottom-mount coordinates" in text
    assert "Conceptual wall-mount coordinates" in text
    assert "Bottom face, Z = 0" in text
    assert "wall-mount height" in text

    movement_table = next(
        table
        for table in doc.tables
        if table.rows[0].cells[0].text == "Axis"
        and table.rows[0].cells[1].text == "Peak movement (mm)"
    )
    movements = [
        float(row.cells[1].text)
        for row in movement_table.rows[1:]
    ]
    assert all(value > 0 for value in movements)


def test_supplier_pack_bytes_are_a_readable_word_document():
    data = supplier_enquiry_pack_bytes(
        _selection_snapshot(),
        layout_box_mm=(600.0, 800.0, 1200.0),
    )

    assert data.startswith(b"PK")
    reopened = Document(BytesIO(data))
    assert "WIRE-ROPE ISOLATOR" in _document_text(reopened)


def test_every_supplier_pack_table_has_fixed_9360_dxa_geometry():
    doc = generate_supplier_enquiry_pack(
        _selection_snapshot(),
        layout_box_mm=(600.0, 800.0, 1200.0),
    )

    assert doc.settings.odd_and_even_pages_header_footer is True
    for header in (
        doc.sections[0].header,
        doc.sections[0].even_page_header,
        doc.sections[0].first_page_header,
    ):
        assert header.paragraphs[0].text == ""
    for footer in (
        doc.sections[0].footer,
        doc.sections[0].even_page_footer,
        doc.sections[0].first_page_footer,
    ):
        assert footer.paragraphs[0]._p.xpath(".//w:fldChar")
    assert doc.tables
    for table in doc.tables:
        table_width = table._tbl.tblPr.first_child_found_in("w:tblW")
        assert table_width.get(qn("w:type")) == "dxa"
        assert int(table_width.get(qn("w:w"))) == CONTENT_WIDTH_DXA

        grid_widths = [
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid
        ]
        assert sum(grid_widths) == CONTENT_WIDTH_DXA
        for row in table.rows:
            cell_widths = [
                int(
                    cell._tc.get_or_add_tcPr()
                    .first_child_found_in("w:tcW")
                    .get(qn("w:w"))
                )
                for cell in row.cells
            ]
            assert cell_widths == grid_widths
