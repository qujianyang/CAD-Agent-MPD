"""
SAR appendix generator — produces a Word (.docx) document with Appendices
B, C, D, E matching the layout of the Variant E-2 Road Trial Safety
Assessment Report (R-J7267-TR008).

  Appendix B — Centre of Gravity, Weight Distribution and Axle Loadings
  Appendix C — Longitudinal Slope Stability Analysis
  Appendix D — Lateral Slope Stability Analysis
  Appendix E — Cornering Stability Analysis

Deterministic: every safety factor and load comes from the validated
mobility_engine. The LLM is never in the numeric path. Figures (vehicle
diagrams, photos) are inserted as labelled placeholders for the engineer
to drop in manually.
"""
import math
from datetime import date as _date
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mobility_engine import (
    Vehicle, Aero, G,
    axle_loads, slope_stability, side_slope_stability, cornering_stability,
)
from mobility_import import MeasurementData, PlatformCG


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def _caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.bold = True
    r.font.size = Pt(10)
    return p


def _figure_placeholder(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[ {text} — insert figure ]")
    r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p


def _formula(doc, text, *, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.bold = bold
    return p


def _grid_table(doc, data, *, header=True, widths=None):
    """data = list of rows; each row = list of cell strings."""
    rows = len(data)
    cols = max(len(r) for r in data)
    t = doc.add_table(rows=rows, cols=cols)
    t.style = "Table Grid"
    for ri, row in enumerate(data):
        for ci in range(cols):
            cell = t.rows[ri].cells[ci]
            txt = str(row[ci]) if ci < len(row) else ""
            cell.text = ""
            run = cell.paragraphs[0].add_run(txt)
            if header and ri == 0:
                run.bold = True
            run.font.size = Pt(9)
    return t


# ---------------------------------------------------------------------------
# Appendix B — CG, Weight Distribution, Axle Loadings
# ---------------------------------------------------------------------------

def _wheel_axle_tables(doc, m: MeasurementData):
    """Wheel-load reading table + axle summary table (shared by B.1 and B.2)."""
    rows = [["Axle Loading Measurement On Flat Ground", "1st Reading (kg)",
             "2nd Reading (kg)", "Average Reading (kg)", "Wheel Diff. (%)"]]
    diff_map = {1: f"{m.front_diff_pct:.1f}", 3: f"{m.rear_diff_pct:.1f}"}
    for i, w in enumerate(m.wheels):
        rows.append([w.label, f"{w.r1:.0f}", f"{w.r2:.0f}", f"{w.avg:.0f}",
                     diff_map.get(i, "")])
    wlabel = "Gross Weight, GW" if m.weight_label == "GW" else "Unladen Transporter Weight, CW"
    rows.append([f"{wlabel} (FL + FR + RL + RR)", "", "", f"{m.gw_kg:.0f}", ""])
    _grid_table(doc, rows)
    doc.add_paragraph()

    ax_rows = [["", "Weight (kg)"],
               ["Front Axle (FL + FR)", f"{m.front_axle_kg:.0f}"],
               ["Rear Axle (RL + RR)", f"{m.rear_axle_kg:.0f}"],
               ["Driver Side (FR + RR)", f"{m.driver_kg:.0f}"],
               ["Kerb Side (FL + RL)", f"{m.kerb_kg:.0f}"]]
    _grid_table(doc, ax_rows)
    doc.add_paragraph()


def _cg_derivation(doc, m: MeasurementData, *, sym, fig_x, fig_y, fig_z, fig_summary,
                   summary_label, summary_caption, datum_note):
    """Xcg/Ycg/Zcg worked derivation + summary CG table (shared by B.1 and B.2)."""
    wl = m.weight_label
    # --- X ---
    doc.add_paragraph().add_run(f"Measured X{sym}").bold = True
    _figure_placeholder(doc, fig_x)
    doc.add_paragraph(f"Taking moments about front axle, distance between centre of gravity "
                      f"and front axle, X{sym}:")
    _formula(doc, f"X{sym} = (RA / {wl}) x WB")
    _formula(doc, f"    = ({m.rear_axle_kg:.0f} / {m.gw_kg:.0f}) x {m.wheelbase_mm:.0f}")
    _formula(doc, f"    = {m.xcg_mm:.1f} mm", bold=True)
    # --- Y ---
    doc.add_paragraph().add_run(f"Measured Y{sym}").bold = True
    _figure_placeholder(doc, fig_y)
    ly_kerb = m.track_mm / 2 + m.ycg_mm
    doc.add_paragraph(f"Taking moments about centre of kerbside tyre, position of Y{sym} "
                      f"from kerbside tyre, Ly:")
    _formula(doc, f"Ly = (RGA / {wl}) x TW")
    _formula(doc, f"   = ({m.driver_kg:.0f} / {m.gw_kg:.0f}) x {m.track_mm:.0f}")
    _formula(doc, f"   = {ly_kerb:.1f} mm  (from kerbside tyre)")
    _formula(doc, f"Y{sym} (from centreline, right +) = {m.ycg_mm:.1f} mm", bold=True)
    # --- Z ---
    doc.add_paragraph().add_run(f"Measured Z{sym}").bold = True
    if m.tilt_tests:
        doc.add_paragraph("The Z-axis is determined by lifting the vehicle from the front at "
                          "several inclination angles and reading the rear-axle load.")
        _figure_placeholder(doc, fig_z)
        _formula(doc, "Z = [(FI - FL) x WB] / (W x tan a) + Rstat")
        z_rows = [["Test", "Inclined angle (deg)", "Radius of Wheel (mm)",
                   "Reading, FI (kg)", "Z-Axis (mm)"]]
        for t in m.tilt_tests:
            z_rows.append([f"{t.case}", f"{t.angle_deg:.1f}", f"{t.radius_mm:.0f}",
                           f"{t.fi_kg:.0f}", f"{t.z_mm:.1f}"])
        z_rows.append(["", "", "", f"Average Z{sym}", f"{m.avg_z_mm:.1f}"])
        _grid_table(doc, z_rows)
    else:
        _formula(doc, f"Z{sym} = {m.zcg_mm:.0f} mm from ground"
                      + (f"  ({m.z_note})" if m.z_note else ""), bold=True)
    doc.add_paragraph()
    # --- summary CG ---
    _figure_placeholder(doc, fig_summary)
    doc.add_paragraph(datum_note)
    cg_rows = [
        [summary_label, "Weight (kg)", "X (mm)", "Y (mm)", "Z (mm)"],
        [m.name, f"{m.gw_kg:.0f}", f"{m.xcg_mm:.1f}", f"{m.ycg_mm:.1f}", f"{m.zcg_mm:.1f}"],
    ]
    _grid_table(doc, cg_rows)
    _caption(doc, summary_caption)


def appendix_b(doc, v: Vehicle, m: MeasurementData,
               m_unladen: MeasurementData = None, shelter: PlatformCG = None):
    doc.add_heading("Appendix B  Centre of Gravity, Weight Distribution and Axle Loadings", level=1)

    # ---- B.1 Integrated Platform (laden) ----
    doc.add_heading("B.1  Integrated Platform", level=2)
    doc.add_heading("B.1.1  Integrated Platform Measured Wheel Load Readings", level=3)
    _wheel_axle_tables(doc, m)
    doc.add_paragraph("Summary of the Axle Loadings:")
    sum_rows = [
        ["State", "Front Axle (kg)", "Rear Axle (kg)", "Total Weight (kg)"],
        ["Laden",
         f"{m.front_axle_kg:.0f} < {v.front_axle_limit_kg:.0f}",
         f"{m.rear_axle_kg:.0f} < {v.rear_axle_limit_kg:.0f}",
         f"{m.gw_kg:.0f} < {v.gvw_limit_kg:.0f}"],
    ]
    _grid_table(doc, sum_rows)
    _caption(doc, "Table B-1: Summary of the Measured Integrated Platform Weight and Axle Loadings")
    doc.add_heading("B.1.2  Determine Centre of Gravity of Integrated Platform", level=3)
    _cg_derivation(doc, m, sym="cg",
                   fig_x="Figure B-1: Measured Xcg", fig_y="Figure B-2: Measured Ycg",
                   fig_z="Figure B-3: Z-axis Measurement by Lifting Front Axle",
                   fig_summary="Figure B-4: Integrated Platform Centre of Gravity",
                   summary_label="Overall CG of Laden Vehicle",
                   summary_caption="Table B-2: Integrated Platform Summary CG Table",
                   datum_note="Note: Datum at X: Vehicle front axle, Y: Centre Axis Right (+), "
                              "Z: Ground level")

    # ---- B.2 Unladen AFE Transporter L1 ----
    if m_unladen is not None:
        doc.add_heading("B.2  Unladen AFE Transporter L1", level=2)
        doc.add_heading("B.2.1  Unladen AFE Transporter L1 Measured Wheel Load Readings", level=3)
        _wheel_axle_tables(doc, m_unladen)
        _cg_derivation(doc, m_unladen, sym="CW",
                       fig_x="Figure B-5: Measured XCW", fig_y="Figure B-6: Measured YCW",
                       fig_z="", fig_summary="Figure B-7: Unladen Transporter Centre of Gravity",
                       summary_label="Overall CG of Unladen Transporter L1",
                       summary_caption="Table B-3: Overall Weight and CG of Unladen AFE Transporter L1",
                       datum_note="Note: Datum at X: Centre of Front Axle, Y: Centre Axis "
                                  "Right (+), Z: Ground Level")

    # ---- B.3 Laden E2 Shelter (derived) ----
    if shelter is not None:
        doc.add_heading("B.3  Determination of Weight and Centre of Gravity of Laden E2 Shelter", level=2)
        doc.add_paragraph(
            "The laden E2 shelter weight and CG are obtained by subtracting the unladen "
            "transporter (CW) from the integrated platform (GW): PL = GW - CW, and "
            "moment balance GW x CGgw = CW x CGcw + PL x CGpl on each axis.")
        cg_rows = [
            ["Overall CG of Laden Shelter", "Weight (kg)", "X (mm)", "Y (mm)", "Z (mm)"],
            [shelter.name, f"{shelter.weight_kg:.0f}", f"{shelter.xcg_mm:.1f}",
             f"{shelter.ycg_mm:.1f}", f"{shelter.zcg_mm:.1f}"],
        ]
        _grid_table(doc, cg_rows)
        doc.add_paragraph("Note: Datum at X: Centre of Front Axle, Y: Centre Axis Right (+), "
                          "Z: Ground Level")
        _caption(doc, "Table B-4: Overall Weight and CG of Laden E2 Shelter")


# ---------------------------------------------------------------------------
# Appendix C — Longitudinal Slope Stability
# ---------------------------------------------------------------------------

# SAR declared angles: the report prints these rounded angles and computes with them.
_SAR_ANGLE = {60: 31.0, 30: 16.7}


def _slope_block(doc, v, grade, direction, fig_label, sec_label, table_label, table_caption):
    r = slope_stability(v, grade, direction, angle_deg=_SAR_ANGLE.get(grade), trig_dp=3)
    angle_disp = int(_SAR_ANGLE.get(grade, round(r.theta_deg)))
    doc.add_heading(sec_label, level=2)
    _figure_placeholder(doc, fig_label)
    if direction == "ascending":
        doc.add_paragraph("Taking Moment about Centre of Rear Axle,")
        lever_txt = "(Wheelbase - Xcg)"
    else:
        doc.add_paragraph("Taking Moment about Centre of Front Axle,")
        lever_txt = "(Xcg)"
    _formula(doc, f"Stabilizing Moment = (mg)(cos{angle_disp}) x [{lever_txt}/1000]")
    _formula(doc, f"    = {G} x {v.gw_kg:.0f} x {math.cos(math.radians(r.theta_deg)):.3f} "
                  f"x ({r.lever_mm:.1f}/1000)")
    _formula(doc, f"    = {r.stab_Nm:,.0f} Nm")
    _formula(doc, f"Overturning Moment = (mg)(sin{angle_disp}) x (Zcg/1000)")
    _formula(doc, f"    = {G} x {v.gw_kg:.0f} x {math.sin(math.radians(r.theta_deg)):.3f} "
                  f"x ({v.zcg_mm:.1f}/1000)")
    _formula(doc, f"    = {r.over_Nm:,.0f} Nm")
    _formula(doc, f"Safety Factor = {r.stab_Nm:,.0f} / {r.over_Nm:,.0f} = {r.SF:.2f}", bold=True)
    sf_rows = [
        [table_label, ""],
        ["Detachments", "Safety Factor"],
        ["Integrated Platform", f"{r.SF:.2f}"],
    ]
    _grid_table(doc, sf_rows, header=False)
    _caption(doc, table_caption)
    return r


def appendix_c(doc, v: Vehicle):
    doc.add_heading("Appendix C  Longitudinal Slope Stability Analysis", level=1)
    _slope_block(doc, v, 60, "ascending",
                 "Figure C-1: 60% (31 deg) Ascending Slope",
                 "C.1  For 60% Ascending Slope Stability",
                 "60% Longitudinal Slope Stability (Ascending Slope)",
                 "Table C-1: 60% Ascending Longitudinal Slope Stability")
    _slope_block(doc, v, 60, "descending",
                 "Figure C-2: 60% (31 deg) Descending Slope",
                 "C.2  For 60% Descending Slope Stability",
                 "60% Longitudinal Slope Stability (Descending Slope)",
                 "Table C-2: 60% Descending Longitudinal Slope Stability")


# ---------------------------------------------------------------------------
# Appendix D — Lateral Slope Stability
# ---------------------------------------------------------------------------

def _side_block(doc, v, grade, side, fig_label, sec_label, table_label, table_caption):
    r = side_slope_stability(v, grade, side, angle_deg=_SAR_ANGLE.get(grade), trig_dp=3)
    angle_disp = f"{_SAR_ANGLE.get(grade, r.theta_deg):.1f}"
    sign = "+" if side == "kerbside" else "-"
    doc.add_heading(sec_label, level=2)
    _figure_placeholder(doc, fig_label)
    _formula(doc, f"Stabilizing Moment = (mg)(cos{angle_disp}) x [(Track Width/2 {sign} Ycg)/1000]")
    _formula(doc, f"    = {G} x {v.gw_kg:.0f} x {math.cos(math.radians(r.theta_deg)):.3f} "
                  f"x ({r.lever_mm:.1f}/1000)")
    _formula(doc, f"    = {r.stab_Nm:,.0f} Nm")
    _formula(doc, f"Overturning Moment = (mg)(sin{angle_disp}) x (Zcg/1000)")
    _formula(doc, f"    = {r.over_Nm:,.0f} Nm")
    _formula(doc, f"Safety Factor = {r.stab_Nm:,.0f} / {r.over_Nm:,.0f} = {r.SF:.2f}", bold=True)
    sf_rows = [
        [table_label, ""],
        ["Detachments", "Safety Factor"],
        ["Integrated Platform", f"{r.SF:.2f}"],
    ]
    _grid_table(doc, sf_rows, header=False)
    _caption(doc, table_caption)
    return r


def appendix_d(doc, v: Vehicle):
    doc.add_heading("Appendix D  Lateral Slope Stability Analysis", level=1)
    _side_block(doc, v, 30, "kerbside",
                "Figure D-1: 30% (16.7 deg) Kerbside Slope",
                "D.1  For 30% Kerbside Slope Stability",
                "30% Lateral Slope Stability (Kerbside)",
                "Table D-1: 30% Kerbside Lateral Slope Stability")
    _side_block(doc, v, 30, "roadside",
                "Figure D-2: 30% (16.7 deg) Roadside Slope",
                "D.2  For 30% Roadside Slope Stability",
                "30% Lateral Slope Stability (Roadside)",
                "Table D-2: 30% Roadside Lateral Slope Stability")


# ---------------------------------------------------------------------------
# Appendix E — Cornering Stability
# ---------------------------------------------------------------------------

def appendix_e(doc, v: Vehicle, aero: Aero,
               speed_kmh=15.0, radius_m=11.0, wind_kmh=60.0, oem_sf=1.5):
    c = cornering_stability(v, aero, speed_kmh, radius_m, wind_kmh)
    S = speed_kmh / 3.6
    Vw = wind_kmh / 3.6

    doc.add_heading("Appendix E  Cornering Stability Analysis", level=1)

    # Wind drag breakdown
    doc.add_paragraph("Wind Drag, FD = 1/2 x CD x Dair x Vair^2 x A")
    wind_rows = [
        ["Symbol", "Description", "Value", "Unit"],
        ["CD", "Drag Coefficient", f"{aero.cd:.1f}", "-"],
        ["Dair", "Density of air", f"{aero.air_density:.2f}", "kg/m3"],
        ["Vair", "Velocity of air (60 km/h)", f"{Vw:.1f}", "m/s"],
        ["A", "Side projected area", f"{aero.side_area_m2:.1f}", "m2"],
        ["FD", "Wind drag force", f"{c.Fw_N:,.0f}", "N"],
    ]
    _grid_table(doc, wind_rows)
    doc.add_paragraph()

    _formula(doc, "Overturning Moment = (m x V^2) x Z/R + FD x h")
    over_rows = [
        ["Symbol", "Description", "Value", "Unit"],
        ["Z", "CG height of the vehicle", f"{v.zcg_mm:.1f}", "mm"],
        ["m", "Weight of the vehicle", f"{v.gw_kg:.0f}", "kg"],
        ["V", f"Cornering speed, {speed_kmh:.0f} km/h", f"{S:.2f}", "m/s"],
        ["R", "Min turning radius", f"{radius_m:.0f}", "m"],
        ["h", "Height where wind force acts", f"{aero.wind_height_m:.2f}", "m"],
    ]
    _grid_table(doc, over_rows)
    _formula(doc, f"    = {c.over_total_Nm:,.0f} Nm")
    doc.add_paragraph()

    _formula(doc, "Stabilizing Moment = m x g x Y'")
    doc.add_paragraph("Y' is the shorter lateral lever (left turn is the worst case, less stable).")
    _formula(doc, f"    = {v.gw_kg:.0f} x {G} x ({c.yprime_mm:.1f}/1000)")
    _formula(doc, f"    = {c.resist_Nm:,.0f} Nm")
    _formula(doc, "Safety Factor = Stabilizing Moment / Overturning Moment")
    _formula(doc, f"    = {c.resist_Nm:,.0f} / {c.over_total_Nm:,.0f} = {c.SF:.2f}", bold=True)

    e_rows = [
        ["Turning Stability", f"Safety Factor of Cornering at {speed_kmh:.0f} km/h"],
        ["Integrated Platform", f"{c.SF:.2f}  (Left turn, worst case)"],
        ["OEM Recommended", f"{oem_sf:.1f}"],
    ]
    _grid_table(doc, e_rows, header=False)
    _caption(doc, "Table E-1: Cornering Stability and Maximum Cornering Speed")
    return c


# ---------------------------------------------------------------------------
# Top-level assembly
# ---------------------------------------------------------------------------

def generate_sar_appendices(
    v: Vehicle,
    m: MeasurementData,
    aero: Optional[Aero] = None,
    *,
    m_unladen: Optional[MeasurementData] = None,
    shelter: Optional[PlatformCG] = None,
    project: str = "Project Spinel",
    variant: str = "Variant E-2",
    out_path: Optional[str] = None,
    speed_kmh: float = 15.0,
    radius_m: float = 11.0,
    wind_kmh: float = 60.0,
) -> Document:
    """Build a .docx containing Appendices B, C, D, E. Saves if out_path given.

    m_unladen / shelter are optional: pass them to include B.2 (unladen
    transporter) and B.3 (laden shelter). Omit for B.1 only.
    """
    aero = aero or Aero()
    doc = Document()

    title = doc.add_heading(f"{project} — {variant}", level=0)
    sub = doc.add_paragraph(
        f"Road Trial Safety Assessment Report — Mobility Appendices (B-E)\n"
        f"Generated {_date.today().isoformat()}  ·  CO-CONFIDENTIAL"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    appendix_b(doc, v, m, m_unladen=m_unladen, shelter=shelter)
    doc.add_page_break()
    appendix_c(doc, v)
    doc.add_page_break()
    appendix_d(doc, v)
    doc.add_page_break()
    appendix_e(doc, v, aero, speed_kmh, radius_m, wind_kmh)

    if out_path:
        doc.save(out_path)
    return doc


if __name__ == "__main__":
    from mobility_import import (
        vehicle_measured, measurement_measured, measurement_unladen, shelter_cg,
    )
    v = vehicle_measured()
    m = measurement_measured()
    mu = measurement_unladen()
    sh = shelter_cg()
    out = "Appendix_BCDE_demo.docx"
    generate_sar_appendices(v, m, m_unladen=mu, shelter=sh, out_path=out)
    print(f"[OK] wrote {out} (B.1 + B.2 + B.3 + C + D + E)")
