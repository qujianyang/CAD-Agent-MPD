"""Generate a deterministic Word supplier enquiry pack for shock isolators."""

from __future__ import annotations

from dataclasses import asdict, is_dataclass
from datetime import date
from io import BytesIO
from typing import Any, Mapping, Optional, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from mount_layout_viz import layout_summary


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
TABLE_CELL_MARGIN_DXA = {
    "top": 80,
    "bottom": 80,
    "start": 120,
    "end": 120,
}
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "666666"
LIGHT_GRAY = "F2F4F7"
RISK_RED = "9B1C1C"
PASS_BLUE = "1F3A5F"


def generate_supplier_enquiry_pack(
    snapshot: Any,
    *,
    project_reference: str = "",
    equipment_reference: str = "",
    prepared_by: str = "",
    supplier: str = "Helical / VMC",
    layout_box_mm: Optional[Sequence[float]] = None,
    clearances_mm: Sequence[Optional[float]] = (None, None, None),
    wall_face: str = "back",
    report_date: Optional[str] = None,
) -> Document:
    """Return a Word supplier enquiry pack from one authoritative snapshot."""
    payload = _snapshot_dict(snapshot)
    if not payload:
        raise ValueError("A shock analysis snapshot is required.")

    analysis_id = str(payload.get("analysis_id") or "Unidentified")
    report_day = report_date or date.today().isoformat()
    project_reference = project_reference.strip() or "To be assigned"
    equipment_reference = equipment_reference.strip() or "Equipment rack / enclosure"
    prepared_by = prepared_by.strip() or "To be completed"
    supplier = supplier.strip() or "To be confirmed"

    doc = Document()
    _configure_document(doc)
    _add_customer_pack_header(
        doc,
        analysis_id=analysis_id,
        project_reference=project_reference,
        equipment_reference=equipment_reference,
        prepared_by=prepared_by,
        supplier=supplier,
        report_day=report_day,
    )

    verdict = str(payload.get("verdict") or "UNKNOWN").upper()
    selected_part = str(payload.get("selected_part") or "No passing part identified")
    _add_status_lead(doc, verdict, selected_part)

    doc.add_heading("1. Purpose and status", level=1)
    doc.add_paragraph(
        "This pack records a preliminary wire-rope isolator selection for "
        "supplier review. The engineering values are inserted directly from "
        "the deterministic Python analysis linked to the analysis ID above."
    )
    _add_key_value_table(
        doc,
        [
            ("Supplier confirmation", "Pending"),
            ("Analysis mode", _display(payload.get("mode"))),
            ("Preliminary verdict", verdict),
            ("Preliminary selection", selected_part),
            ("AI-generated calculation values", "None"),
        ],
    )

    doc.add_heading("2. Equipment and shock requirements", level=1)
    _add_key_value_table(
        doc,
        [
            ("Total assessed mass", _unit(payload.get("mass_kg"), "kg")),
            ("Bottom isolators", _display(payload.get("bottom_mounts"))),
            ("Wall isolators", _display(payload.get("wall_mounts"))),
            ("Input shock", _unit(payload.get("input_shock_g"), "G")),
            ("Pulse duration", _unit(payload.get("pulse_duration_ms"), "ms")),
            ("Pulse profile", _display(payload.get("pulse_shape"))),
            (
                "Maximum transmitted shock",
                _unit(payload.get("transmitted_g_limit"), "G"),
            ),
        ],
    )

    doc.add_heading("3. Proposed mounting arrangement", level=1)
    _add_key_value_table(
        doc,
        [
            (
                "Arrangement",
                f"{payload.get('bottom_mounts', 0)} bottom + "
                f"{payload.get('wall_mounts', 0)} wall isolators",
            ),
            ("Wall mounting face", wall_face),
            ("Equipment envelope", _box_text(layout_box_mm)),
            ("Available clearance X / Y / Z", _clearance_text(clearances_mm)),
        ],
    )
    _add_layout_coordinates(
        doc,
        payload,
        layout_box_mm=layout_box_mm,
        clearances_mm=clearances_mm,
        wall_face=wall_face,
    )
    doc.add_paragraph(
        "Mount coordinates are a conceptual, evenly distributed arrangement "
        "for supplier discussion. Final interfaces, wall-mount height, fasteners, "
        "access and local structure require installation drawing review.",
        style="Supplier Note",
    )

    doc.add_heading("4. Preliminary selection", level=1)
    _add_key_value_table(
        doc,
        [
            ("Recommended part", selected_part),
            ("Series", _display(payload.get("selected_series"))),
            ("Governing load case", _display(payload.get("governing_case"))),
            (
                "Governing constraint",
                _display(payload.get("governing_constraint")),
            ),
            (
                "Governing utilization",
                _unit(payload.get("governing_use_pct"), "%", decimals=0),
            ),
            (
                "Worst transmitted shock",
                _unit(payload.get("worst_transmitted_g"), "G"),
            ),
            (
                "Worst calculated movement",
                _unit(payload.get("worst_movement_mm"), "mm"),
            ),
            (
                "Static load / published rating",
                _static_text(
                    payload.get("static_load_daN"),
                    payload.get("static_rating_daN"),
                ),
            ),
            (
                "Validation level",
                _display(payload.get("validation_level") or "catalog data"),
            ),
        ],
    )

    doc.add_heading("5. Four-case calculation results", level=1)
    _add_load_case_table(doc, payload.get("load_cases") or ())

    doc.add_heading("6. Catalogue alternatives", level=1)
    _add_alternative_table(doc, payload.get("alternatives") or ())

    doc.add_page_break()
    doc.add_heading("7. Warnings, assumptions and limitations", level=1)
    warnings = [str(item) for item in payload.get("warnings") or ()]
    if warnings:
        doc.add_paragraph("Analysis warnings:", style="Supplier Label")
        for warning in warnings:
            _add_list_item(doc, warning, numbered=False)
    else:
        doc.add_paragraph("No calculation warnings were generated.")

    doc.add_paragraph("Model assumptions:", style="Supplier Label")
    assumptions = [
        "The rack is represented as a rigid body with simplified equal load distribution.",
        "Catalogue average stiffness and published travel values are used.",
        "The four load cases follow the validated project calculation workflow.",
        "Installation clearance is a design input, not a measured manufacturing tolerance.",
        "The preliminary calculation does not replace nonlinear vendor analysis or testing.",
    ]
    for assumption in assumptions:
        _add_list_item(doc, assumption, numbered=False)

    doc.add_heading("8. Information requested from supplier", level=1)
    supplier_questions = [
        f"Confirm whether {selected_part} is the recommended isolator for the stated requirements.",
        "Confirm static capacity in each installed orientation.",
        "Provide the latest directional load-deflection and stiffness data.",
        "Confirm the proposed bottom and wall-mount arrangement and orientation.",
        "Confirm available shock travel and minimum installation clearance.",
        "Identify required brackets, fasteners, preload and installation controls.",
        "Provide applicable qualification, test or similarity evidence.",
        "Identify any frequency, temperature, corrosion or service-life limitations.",
    ]
    for question in supplier_questions:
        _add_list_item(doc, question, numbered=True)

    doc.add_heading("9. Supplier response record", level=1)
    _add_key_value_table(
        doc,
        [
            ("Supplier recommendation", "Pending"),
            ("Confirmed part number", "Pending"),
            ("Confirmed arrangement", "Pending"),
            ("Supporting report / drawing", "Pending"),
            ("Differences from preliminary result", "Pending"),
            ("Supplier representative / date", "Pending"),
        ],
    )

    doc.add_page_break()
    doc.add_heading("10. Approval boundary", level=1)
    boundary = doc.add_paragraph()
    boundary.style = doc.styles["Supplier Boundary"]
    boundary.add_run(
        "This document is a preliminary engineering selection package intended "
        "for supplier review. It does not replace supplier confirmation, detailed "
        "installation review, physical qualification testing or project approval."
    )

    doc.add_paragraph()
    signoff = _add_table(
        doc,
        [
            ["Prepared by", "Reviewed by", "Supplier acknowledgement"],
            [prepared_by, "", ""],
            ["Date: " + report_day, "Date:", "Date:"],
        ],
        widths_dxa=[3120, 3120, 3120],
        header=True,
    )
    for row in signoff.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(14)

    return doc


def supplier_enquiry_pack_bytes(snapshot: Any, **kwargs) -> bytes:
    """Return the generated Word document as download-ready bytes."""
    doc = generate_supplier_enquiry_pack(snapshot, **kwargs)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = True

    _set_style(
        doc.styles["Normal"],
        font_name="Calibri",
        size=11,
        color="000000",
        before=0,
        after=6,
        line_spacing=1.10,
    )
    _set_style(
        doc.styles["Heading 1"],
        font_name="Calibri",
        size=16,
        color=BLUE,
        bold=True,
        before=16,
        after=8,
        line_spacing=1.0,
    )
    _set_style(
        doc.styles["Heading 2"],
        font_name="Calibri",
        size=13,
        color=BLUE,
        bold=True,
        before=12,
        after=6,
        line_spacing=1.0,
    )
    _set_style(
        doc.styles["Heading 3"],
        font_name="Calibri",
        size=12,
        color=DARK_BLUE,
        bold=True,
        before=8,
        after=4,
        line_spacing=1.0,
    )
    _set_style(
        doc.styles["List Bullet"],
        font_name="Calibri",
        size=11,
        color="000000",
        before=0,
        after=8,
        line_spacing=1.167,
    )
    _set_style(
        doc.styles["List Number"],
        font_name="Calibri",
        size=11,
        color="000000",
        before=0,
        after=8,
        line_spacing=1.167,
    )

    note_style = _get_or_add_style(doc, "Supplier Note")
    _set_style(
        note_style,
        font_name="Calibri",
        size=10,
        color=MUTED,
        italic=True,
        before=4,
        after=8,
        line_spacing=1.10,
    )
    label_style = _get_or_add_style(doc, "Supplier Label")
    _set_style(
        label_style,
        font_name="Calibri",
        size=11,
        color=DARK_BLUE,
        bold=True,
        before=6,
        after=4,
        line_spacing=1.10,
    )
    boundary_style = _get_or_add_style(doc, "Supplier Boundary")
    _set_style(
        boundary_style,
        font_name="Calibri",
        size=10.5,
        color=RISK_RED,
        bold=True,
        before=8,
        after=10,
        line_spacing=1.10,
    )

    for header_part in (
        section.header,
        section.even_page_header,
        section.first_page_header,
    ):
        header = header_part.paragraphs[0]
        header.text = ""
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.left_indent = Pt(0)
        header.paragraph_format.right_indent = Pt(0)
        header.paragraph_format.first_line_indent = Pt(0)
        header.paragraph_format.tab_stops.clear_all()
        header.paragraph_format.space_after = Pt(0)

    for footer_part in (
        section.footer,
        section.even_page_footer,
        section.first_page_footer,
    ):
        footer = footer_part.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.left_indent = Pt(0)
        footer.paragraph_format.right_indent = Pt(0)
        footer.paragraph_format.first_line_indent = Pt(0)
        footer.paragraph_format.tab_stops.clear_all()
        _add_page_field(footer)


def _add_customer_pack_header(
    doc: Document,
    *,
    analysis_id: str,
    project_reference: str,
    equipment_reference: str,
    prepared_by: str,
    supplier: str,
    report_day: str,
) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("WIRE-ROPE ISOLATOR")
    _format_run(title_run, size=23, color="000000", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("Preliminary Supplier Enquiry Pack")
    _format_run(subtitle_run, size=14, color=MUTED, bold=True)

    _add_key_value_table(
        doc,
        [
            ("Project / enquiry reference", project_reference),
            ("Equipment reference", equipment_reference),
            ("Analysis ID", analysis_id),
            ("Prepared by", prepared_by),
            ("Date", report_day),
            ("Intended supplier", supplier),
        ],
    )


def _add_status_lead(doc: Document, verdict: str, selected_part: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    label = paragraph.add_run(f"PRELIMINARY {verdict}")
    _format_run(
        label,
        size=12,
        color=PASS_BLUE if verdict == "PASS" else RISK_RED,
        bold=True,
    )
    detail = paragraph.add_run(f"    {selected_part}")
    _format_run(detail, size=12, color="000000", bold=True)


def _add_load_case_table(doc: Document, load_cases: Sequence[Mapping]) -> None:
    rows = [
        [
            "Load case",
            "Status",
            "Load\nkg",
            "fn\nHz",
            "GT / limit\nG",
            "Movement / limit\nmm",
            "Impulse model",
        ]
    ]
    for case in load_cases:
        rows.append(
            [
                _display(case.get("name")),
                _display(case.get("status")),
                _number(case.get("load_kg")),
                _number(case.get("natural_frequency_hz")),
                f"{_number(case.get('transmitted_g'))} / "
                f"{_number(case.get('transmitted_g_limit'))}",
                f"{_number(case.get('movement_mm'))} / "
                f"{_number(case.get('movement_limit_mm'))}",
                "Within range" if case.get("impulse_model_valid") else "Review",
            ]
        )
    _add_table(
        doc,
        rows,
        widths_dxa=[1950, 850, 850, 850, 1450, 1650, 1760],
        header=True,
        center_columns={1, 2, 3, 4, 5, 6},
    )


def _add_alternative_table(doc: Document, alternatives: Sequence[Mapping]) -> None:
    if not alternatives:
        doc.add_paragraph(
            "No catalogue alternatives are recorded for this analysis mode."
        )
        return
    rows = [
        [
            "Part",
            "Series",
            "Status",
            "Worst GT use",
            "Worst movement use",
            "Worst overall use",
        ]
    ]
    for item in alternatives:
        rows.append(
            [
                _display(item.get("part_number")),
                _display(item.get("series")),
                _display(item.get("status")),
                _unit(item.get("worst_transmitted_g_use_pct"), "%", decimals=0),
                _unit(item.get("worst_movement_use_pct"), "%", decimals=0),
                _unit(item.get("worst_overall_use_pct"), "%", decimals=0),
            ]
        )
    _add_table(
        doc,
        rows,
        widths_dxa=[1700, 1200, 900, 1750, 1900, 1910],
        header=True,
        center_columns={1, 2, 3, 4, 5},
    )


def _add_layout_coordinates(
    doc: Document,
    payload: Mapping[str, Any],
    *,
    layout_box_mm: Optional[Sequence[float]],
    clearances_mm: Sequence[Optional[float]],
    wall_face: str,
) -> None:
    if not layout_box_mm or len(layout_box_mm) != 3:
        doc.add_paragraph(
            "Mount coordinates are not included because enclosure dimensions "
            "were not supplied."
        )
        return

    width, depth, height = (float(value) for value in layout_box_mm)
    deflections = _axis_deflections(payload.get("load_cases") or ())
    layout = layout_summary(
        width,
        depth,
        height,
        int(payload.get("bottom_mounts") or 0),
        int(payload.get("wall_mounts") or 0),
        deflections=deflections,
        clearances=clearances_mm,
        wall_face=wall_face,
    )

    doc.add_heading("3.1 Conceptual bottom-mount coordinates", level=2)
    bottom_rows = [["Mount", "X (mm)", "Y (mm)", "Interface"]]
    for index, (x, y) in enumerate(layout["bottom_positions_xy"], 1):
        bottom_rows.append([index, _number(x), _number(y), "Bottom face, Z = 0"])
    _add_table(
        doc,
        bottom_rows,
        widths_dxa=[900, 1600, 1600, 5260],
        header=True,
        center_columns={0, 1, 2},
    )

    doc.add_heading("3.2 Conceptual wall-mount coordinates", level=2)
    wall_rows = [["Mount", "X (mm)", "Y (mm)", "Z (mm)", "Face"]]
    for index, (x, y, z) in enumerate(layout["wall_positions_xyz"], 1):
        wall_rows.append(
            [index, _number(x), _number(y), _number(z), wall_face]
        )
    _add_table(
        doc,
        wall_rows,
        widths_dxa=[900, 1500, 1500, 1500, 3960],
        header=True,
        center_columns={0, 1, 2, 3},
    )

    doc.add_heading("3.3 Calculated movement and clearance", level=2)
    movement_rows = [["Axis", "Peak movement (mm)", "Clearance (mm)", "Status"]]
    for axis in "XYZ":
        item = layout["clearance"][axis]
        clearance = (
            "Not specified"
            if item["clearance_mm"] is None
            else _number(item["clearance_mm"])
        )
        status = (
            "INTERFERENCE"
            if item["interference"]
            else "Not checked"
            if item["clearance_mm"] is None
            else "Clear"
        )
        movement_rows.append(
            [axis, _number(item["sway_mm"]), clearance, status]
        )
    _add_table(
        doc,
        movement_rows,
        widths_dxa=[1000, 2600, 2600, 3160],
        header=True,
        center_columns={0, 1, 2, 3},
    )


def _axis_deflections(load_cases: Sequence[Mapping]) -> tuple[float, float, float]:
    by_name = {
        str(case.get("name") or "").lower(): float(case.get("movement_mm") or 0.0)
        for case in load_cases
    }

    def _find(*term_groups: tuple[str, ...]) -> float:
        for name, movement in by_name.items():
            for terms in term_groups:
                if all(term in name for term in terms):
                    return movement
        return 0.0

    comp_bottom = _find(("comp", "bottom"), ("bottom", "compression"))
    comp_wall = _find(("comp", "wall"), ("wall", "compression"))
    roll_wall = _find(("roll", "wall"), ("wall", "shear"))
    roll_bottom = _find(("roll", "bottom"), ("bottom", "shear"))
    return (
        max(roll_wall, roll_bottom),
        max(comp_wall, roll_bottom),
        max(comp_bottom, roll_wall),
    )


def _add_key_value_table(doc: Document, rows: Sequence[tuple[str, str]]):
    return _add_table(
        doc,
        [[label, value] for label, value in rows],
        widths_dxa=[2600, 6760],
        header=False,
        label_column=0,
    )


def _add_table(
    doc: Document,
    rows: Sequence[Sequence[Any]],
    *,
    widths_dxa: Sequence[int],
    header: bool,
    center_columns: Optional[set[int]] = None,
    label_column: Optional[int] = None,
):
    if not rows:
        raise ValueError("Table rows cannot be empty.")
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("Table column widths must total 9360 DXA.")
    column_count = len(widths_dxa)
    if any(len(row) != column_count for row in rows):
        raise ValueError("Every table row must match the column count.")

    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.autofit = False
    _set_table_geometry(table, widths_dxa)
    center_columns = center_columns or set()

    for row_index, row_values in enumerate(rows):
        for column_index, value in enumerate(row_values):
            cell = table.rows[row_index].cells[column_index]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if column_index in center_columns:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            _format_run(run, size=9, color="000000")
            if header and row_index == 0:
                run.bold = True
                _shade_cell(cell, LIGHT_GRAY)
            elif label_column == column_index:
                run.bold = True
                _shade_cell(cell, LIGHT_GRAY)
        _prevent_row_split(table.rows[row_index])
    if header:
        _repeat_header_row(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def _set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    table_pr = table._tbl.tblPr
    table_width = table_pr.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(CONTENT_WIDTH_DXA))

    table_indent = table_pr.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        table_pr.append(table_indent)
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    layout = table_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    _set_table_cell_margins(table_pr)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        grid.append(grid_column)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440.0)
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:type"), "dxa")
            cell_width.set(qn("w:w"), str(width))


def _set_table_cell_margins(table_pr) -> None:
    margins = table_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        table_pr.append(margins)
    for edge, width in TABLE_CELL_MARGIN_DXA.items():
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _repeat_header_row(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def _prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    row_properties.append(cannot_split)


def _add_list_item(doc: Document, text: str, *, numbered: bool) -> None:
    style_name = "List Number" if numbered else "List Bullet"
    paragraph = doc.add_paragraph(text, style=style_name)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)


def _get_or_add_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _set_style(
    style,
    *,
    font_name: str,
    size: float,
    color: str,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    before: float,
    after: float,
    line_spacing: float,
) -> None:
    style.font.name = font_name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font_name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font_name)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line_spacing
    paragraph_format.keep_with_next = style.name.startswith("Heading")


def _format_run(
    run,
    *,
    size: float,
    color: str,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)
    _format_run(run, size=9, color=MUTED)


def _snapshot_dict(snapshot: Any) -> dict[str, Any]:
    if isinstance(snapshot, Mapping):
        return dict(snapshot)
    if hasattr(snapshot, "to_dict"):
        return dict(snapshot.to_dict())
    if is_dataclass(snapshot):
        return asdict(snapshot)
    return {}


def _display(value: Any) -> str:
    if value is None or value == "":
        return "Not available"
    return str(value)


def _number(value: Any, decimals: int = 2) -> str:
    if value is None:
        return "N/A"
    number = float(value)
    text = f"{number:.{decimals}f}"
    return text.rstrip("0").rstrip(".")


def _unit(value: Any, unit: str, decimals: int = 2) -> str:
    if value is None:
        return "Not available"
    separator = "" if unit == "%" else " "
    return f"{_number(value, decimals)}{separator}{unit}"


def _static_text(load_daN: Any, rating_daN: Any) -> str:
    if load_daN is None:
        return "Not available"
    if rating_daN is None:
        return f"{_number(load_daN)} daN / vendor confirmation required"
    return f"{_number(load_daN)} / {_number(rating_daN)} daN"


def _box_text(layout_box_mm: Optional[Sequence[float]]) -> str:
    if not layout_box_mm or len(layout_box_mm) != 3:
        return "Not supplied"
    width, depth, height = layout_box_mm
    return (
        f"{_number(width)} W x {_number(depth)} D x "
        f"{_number(height)} H mm"
    )


def _clearance_text(clearances_mm: Sequence[Optional[float]]) -> str:
    values = []
    for value in clearances_mm:
        if value is None or float(value) <= 0:
            values.append("not specified")
        else:
            values.append(f"{_number(value)} mm")
    return " / ".join(values)
